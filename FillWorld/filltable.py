from docxtpl import DocxTemplate
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


# 替换
def subtag():
    tpl = DocxTemplate('1.docx')
    sd = tpl.new_subdoc()
    sd.add_paragraph('测试成功')
    context = {
        '测试一下': sd,
    }

    tpl.render(context)
    tpl.save('1_new.docx')


# table = sd.add_table(rows=3, cols=3)
#
# cells = table.rows[0].cells
#
# cells[0].text = "Gene"
# cells[1].text = "Drug"
# cells[2].text = "Rank"


def Setcenter():
    docx = Document('1_new.docx')
    tables = docx.tables
    table = tables[0]
    cell = table.cell(0, 0)
    paragraph = cell.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # cell = table.cell(0, 0)
    # print(cell.text)
    # # pa = cell.paragraphs[0]
    # table.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    # for row in range(2):
    #     for col in range(1):
    #         cell = table.cell(row, col)
    #         pa = cell.paragraphs[0]
    #         pa.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

subtag()
Setcenter()
